import random
import pprint
from docx import Document
from docx.shared import Pt, Inches, Cm, Mm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_UNDERLINE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

# +++ ich hoffe, diesen Code muss sich nie jemand angucken, da er absolut schrecklich aussieht und komplett unoptimiert ist +++
# TODO: Code lesbarer, wartbarer und optimierter machen...mach ich eh nicht

def create_docx(plan:list):
    dates, times, event = read_dates()
    document = Document()


    # A4-Format
    section = document.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(25.4)
    section.right_margin = Mm(25.4)
    section.top_margin = Mm(25.4)
    section.bottom_margin = Mm(25.4)
    section.header_distance = Mm(12.7)
    section.footer_distance = Mm(12.7)


    # Benutzerdefinierten Stil für das Dokument definieren
    style = document.styles['Normal']
    font = style.font
    font.name = 'Microsoft Sans Serif'
    font.size = Pt(11)


    # Ueberschrift von Dokument
    heading_text = f"Messdienerplan vom {plan[0][0]} bis zum {plan[-1][0]}\n"
    heading = document.add_paragraph()
    run = heading.add_run(heading_text)
    font = run.font
    font.size = Pt(14)
    run.underline = True  # Unterstrichen
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Mittelbuendig ausrichten


    # Plan als Tabelle
    table_header = ["Datum", "Uhrzeit", "Anlass", "Messdiener"]
    rows:int = (len(plan) + 1) # Zeilen + 1, da Tabellenueberschrift
    cols = 4 # Spalten
    table = document.add_table(rows, cols) # neue Tabelle
    table.style = "Table Grid"
    for headptr in range(0,4): # Tabellenueberschriften
        table.rows[0].cells[headptr].text = table_header[headptr]    
    for dateidx in range(0,len(plan)):
        table.rows[dateidx+1].cells[0].text = plan[dateidx][0] # Datum in erste Spalte
        table.rows[dateidx+1].cells[1].text = times[plan[dateidx][0]] + " Uhr" # Uhrzeit in zweite Spalte
        table.rows[dateidx+1].cells[2].text = event[plan[dateidx][0]] # Anlass in dritte Spalte
        table.rows[dateidx+1].cells[3].text = formatiere_daten(plan[dateidx][1]) # Eingeteilte Personen in vierte Spalte

    # Breite der Spalten
    for cell in table.columns[0].cells: # Datum
        cell.width = Inches(2)
    for cell in table.columns[1].cells: # Uhrzeit
        cell.width = Inches(1.5)
    for cell in table.columns[2].cells: # Anlass
        cell.width = Inches(2)
    for cell in table.columns[3].cells: # Personen
        cell.width = Inches(6)    

    # Inhalt der Zellen mittig ausrichten
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER



    # Dokument speichern (ersetzt vorhandenes Dokument)
    try:
        document.save("Plan.docx")
    except PermissionError:
        print("##### Word-Dokument noch offen, Schreiben nicht moeglich #####")

def formatiere_daten(daten): # formatiert Namen richtig mit + zwischen Vornamen und , zwischen Namen
    formatierte_daten = []
    for nachname, vornamen in daten:
        formatierter_name = ' + '.join(vornamen) + " " + nachname
        formatierte_daten.append(formatierter_name)
    return ', '.join(formatierte_daten)

def get_dist(name:str, how_often_so_far:dict, to_date_idx, plan:list) -> int: # gibt Distanz einer Person zu seiner letzten Einteilung zurueck (99999 wenn bisher noch nicht dran)
    if how_often_so_far[name] == 0: # wenn name bisher noch nicht eingeteilt
        return 99999
    else:
        lower_limit = 99999 # Distanz zu: vor uebergebenem Termin eingeteilt
        higher_limit = 99999 # Distanz zu: nach uebergebenem Termin eingeteilt
        for i in range(to_date_idx, -1, -1): # finde lower_limit
            if plan[i][1]: # wenn min. 1 Slot belegt
                for k in range(0,len(plan[i][1])): # geht alle bisher eingeteilten Personen durch
                    if name in plan[i][1][k]: # name gefunden
                        lower_limit = to_date_idx - i # berechnet Differenz
        for i in range(to_date_idx,len(plan)): # finde higher_limit
            if plan[i][1]:
                for k in range(0,len(plan[i][1])):
                    if name in plan[i][1][k]:
                        higher_limit = i - to_date_idx
        return min(lower_limit, higher_limit)

def read_persons() -> list:
    with open("personenliste.txt", newline="", encoding="utf-8") as persons_txt:
        persons = []
        for line in persons_txt:
            full_name:list = line.split(" ") # teilt Namen in Vorname und Nachname
            full_name[-1] = full_name[-1].replace("\r\n", "") # entfernt alle \r\n
            first_name = full_name[0]
            last_name = full_name[1]
            found_last_name:bool = False # Indikator ob Nachname schon drin ist
            for entry in persons: # alle bisherigen Eintraege
                if last_name in entry: # Nachname schon vorhanden
                    if not (first_name in entry[1]): # Vorname noch nicht vorhanden
                        entry[1].append(first_name) # fuegt Vorname hinzu
                        found_last_name = True # Nachname nicht nochmal hinzufuegen
            if not found_last_name: # Nachnamenoch nicht vorhanden
                persons.append([last_name, [first_name]]) # fuegt ganzen Namen als neue Familie hinzu
    return persons

def read_dates(): # -> dates:list, times:dict, event:dict
    with open("daten.txt", newline="", encoding="utf-8") as dates_txt:
        dates = []
        times = {}
        events = {}

        for line in dates_txt:
            if len(line) > 2 or not ("\r\n" in line) : # wenn Zeile nicht leer
                line_data = line.split("\t") # teilt Daten an Tabstopp
                line_data[-1] = line_data[-1].replace("\r\n", "") # entfernt alle \r\n
                for idx in range(0,len(line_data)): # wertet alle absichtlichen \n auch als \n aus
                    line_data[idx] = line_data[idx].replace("\\n", "\n")
                current_date = line_data[0]
                current_time = line_data[1]
                current_event = line_data[2]
                current_slots = int(line_data[3])

                # Listen und Dictionaries fuellen
                dates.append([current_date, current_slots])
                times[current_date] = current_time
                events[current_date] = current_event
    return dates, times, events

def read_unavailable_dates() -> dict:
    with open("terminausnahmen.txt", newline="", encoding="utf-8") as unavailable_dates_txt:
        unavailable_dates:dict = {}
        persons = read_persons()
        for person in persons: # init leere liste fuer alle abwesenden Daten einer Person
            unavailable_dates[person[0]] = []
        for line in unavailable_dates_txt:
            if len(line) > 2 or not ("\r\n" in line) : # wenn Zeile nicht leer
                line_data = line.split("\t") # trennt nachname und Daten
                line_data[-1] = line_data[-1].replace("\r\n", "") # entfernt alle \r\n
                last_name = line_data[0]
                dates = line_data[1].split(", ") # trennt Daten
                for date in dates: # fuegt Daten in unavailable_dates der Person hinzu
                    unavailable_dates[last_name].append(date)
    return unavailable_dates

def generate_plan():
    persons:list = read_persons()
    dates, times_dummy, events_dummy = read_dates()
    unavailable_dates = read_unavailable_dates()
    plan = []
    random.shuffle(persons) # zu Beginn alle Personen mischen, damit nicht immer gleiche mit gleichen zusammen
    number_of_slots:int = sum(date[1] for date in dates) # insgesamt zu verteilende Einteilungen
    number_of_persons:int = sum(len(person[1]) for person in persons) # Anzahl Personen
    average:int = number_of_slots // number_of_persons # durchschnittliche, abgerundete anzahl pro Person
    lower_limit = average
    how_often_so_far = {name[0]: 0 for name in persons} # neues Dictionary mit key=nachname und value=0, da jeder am Anfang noch nicht eingeteilt ist

    only_dates = [] # alle daten ohne anzahl slots
    for date in dates:
        only_dates.append(date[0])

    available_dates = {}
    for last_name in persons:
        all_dates = only_dates.copy() # kopie von only_dates
        for date in unavailable_dates[last_name[0]]:
            if date in all_dates: # valides Datum, welches angegeben wurde
                all_dates.remove(date)
            else:
                print("angegebenes Datum ", date, " von ", last_name[0], 
                      " wurde falsch geschrieben oder hat ungueltiges Format. Bitte pruefen, da Person sonst evtl dort eingeteilt wurde")
        available_dates[last_name[0]] = all_dates # fuegt nachnamen und verfuegbare daten in die neue liste

    for date in only_dates: # erstellt Plan mit allen Daten und leeren Slots
        plan.append([date,[]])

    can_only_few_times = [] # alle die nur lower_limit oder weniger oft koennen
    for person in persons: # fuer jeden Nachnamen in persons
        if unavailable_dates[person[0]]: # wenn person daten angegeben hat, an denen er nicht kann
            if (len(available_dates[person[0]])) <= lower_limit: # wenn nachname unterdurchschnittlich oft kann
                can_only_few_times.append(person)
    
    for last_name in can_only_few_times: # alle die ganz wenig koennen als erstes in Plan verteilen
        for available_date in available_dates[last_name[0]]: # fuer alle daten an denen last_name kann
            if dates[only_dates.index(available_date) ][1] > 0: # wenn noch Slots frei sind
                idx = only_dates.index(available_date)
                plan[idx][1].append(last_name)
                how_often_so_far[last_name[0]] += 1
                dates[idx][1] -= len(last_name[1])
    
    for person in can_only_few_times:            
        available_dates.pop(person[0],None) # Personen, die wenig koennen, werden vorher verteilt und sollen nachher nicht mehr beruecksichtigt werden
        unavailable_dates.pop(person[0],None)

    
    for date_idx in range(0,len(dates)):
        required:int = dates[date_idx][1] # Anzahl benoetigte Personen fuer diesen Termin
        can_on_date:list = []
        for name_key in available_dates: # name_key ist nur nachname
            if only_dates[date_idx] in available_dates[name_key]: # wenn Person an datum kann
                can_on_date.append(name_key)
        still_required = required
        while not still_required < 1: # zaehlt von required bis 1 runter
            for last_name in can_on_date: # loescht alle geschwister, die zu viel fuer den tag waeren
                for person in persons: # durchsucht alle personen
                    if person[0] == last_name and len(person[1]) > still_required: # wenn anzahl geschwister zu viel waeren
                        can_on_date.remove(last_name) # loescht nachname aus liste

            if (len(can_on_date) > 1): # wenn mehr als 1 Person kann, filtere weiter
                least_used_persons = []
                min_usage:int = min(how_often_so_far[name] for name in can_on_date) # minimale Anzahl bisher dran
                least_used_persons.extend([name for name in can_on_date if how_often_so_far[name] == min_usage]) # von denen: alle die am wenigsten dran waren
                filtered_persons = least_used_persons

                max_dist = max(get_dist(name, how_often_so_far, date_idx, plan) for name in least_used_persons) # Personen, die am laengsten nicht mehr dran waren
                filtered_persons = [name for name in least_used_persons if get_dist(name, how_often_so_far, date_idx, plan) == max_dist]

                # if (len(filtered_persons) > 1):
                #     max_siblings = 0
                #     max_siblings_name = "last_name_placeholder"
                #     for last_name in filtered_persons: # sucht die Familie mit den meisten Geschwistern, damit grosse Gruppen zuerst eingeteilt werden
                #         for person in persons:
                #             if person[0] == last_name and (len(person[1]) > max_siblings):
                #                 max_siblings = len(person[1])
                #                 max_siblings_name = last_name
                #     filtered_persons = [max_siblings_name] # falls mehrere Geschwister gleiche Werte haben, wird Nachname weiter vorne in Liste genommen

                if (len(filtered_persons) > 1): # wenn immer noch mehrere Personen, dann alle mit wenigsten verfuegbaren Terminen
                    min_dates = 99999
                    min_dates_name = "last_name_placeholder"
                    for last_name in filtered_persons:
                        if len(available_dates[last_name]) < min_dates:
                            min_dates = len(available_dates[last_name])
                            min_dates_name = last_name
                    filtered_persons = [min_dates_name]

                
            elif (len(can_on_date) == 0):
                print("Fuer den ", only_dates[date_idx], " koennen zu wenige Personen")
                filtered_persons = []
                break
            else:
                filtered_persons = can_on_date
            if not len(filtered_persons) == 0: # wenn nach allen Filtern immernoch Personen drin sind (nur zur Absicherung)
                random_person_idx = random.randint(0,len(filtered_persons)-1) # random Person aus verbleibenden auswaehlen
                for person in persons: # sucht ausgesuchte Person(en) in persons
                    if person[0] == filtered_persons[random_person_idx]: # Person(en) gefunden
                        plan[date_idx][1].append(person) # erweitert Namen bei Termin (fuellt Slots)
                        still_required -= len(person[1])
                how_often_so_far[filtered_persons[random_person_idx]] += 1
                can_on_date.remove(filtered_persons[random_person_idx])

    return plan, how_often_so_far


if __name__ == "__main__":
    print("===== Dienstplan =====")

    try:
        plan, how_often = generate_plan()
        # pprint.pprint(plan, width=30)
        print("======================")
        pprint.pprint(how_often)

        create_docx(plan)

        print("\nStrg+C zum Beenden")
        while True:
            pass
    except Exception as e:
        print("Fehler aufgetreten: ", e)
        print("Strg+C zum Beenden")
        while True:
            pass
        

            
